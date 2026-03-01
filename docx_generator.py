import re
import io
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def setup_styles(doc: Document):
    """Configure document styles for consistent formatting."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Title style
    if 'IR Title' not in [s.name for s in doc.styles]:
        title_style = doc.styles.add_style('IR Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(22)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        title_style.paragraph_format.space_after = Pt(4)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle style
    if 'IR Subtitle' not in [s.name for s in doc.styles]:
        sub_style = doc.styles.add_style('IR Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        sub_style.font.name = 'Calibri'
        sub_style.font.size = Pt(12)
        sub_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        sub_style.paragraph_format.space_after = Pt(12)

    # Section heading style
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    # Subsection heading style
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    # Step heading style
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    # Code block style
    if 'Code Block' not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)
        code_style.paragraph_format.left_indent = Inches(0.3)

    # Warning/log gap style
    if 'Warning Text' not in [s.name for s in doc.styles]:
        warn_style = doc.styles.add_style('Warning Text', WD_STYLE_TYPE.PARAGRAPH)
        warn_style.font.name = 'Calibri'
        warn_style.font.size = Pt(11)
        warn_style.font.color.rgb = RGBColor(0xB3, 0x5A, 0x00)
        warn_style.paragraph_format.space_before = Pt(6)
        warn_style.paragraph_format.space_after = Pt(6)
        warn_style.paragraph_format.left_indent = Inches(0.2)

    # Team heading style
    if 'Team Heading' not in [s.name for s in doc.styles]:
        team_style = doc.styles.add_style('Team Heading', WD_STYLE_TYPE.PARAGRAPH)
        team_style.font.name = 'Calibri'
        team_style.font.size = Pt(12)
        team_style.font.bold = True
        team_style.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        team_style.paragraph_format.space_before = Pt(10)
        team_style.paragraph_format.space_after = Pt(4)


def add_code_block(doc: Document, code_text: str):
    """Add a formatted code block with dark background."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['Code Block']
        # Add shading to paragraph
        pPr = p._p.get_or_add_pPr()
        shd = pPr.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '1E1E1E'
        })
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)


def add_metadata_field(doc: Document, label: str, value: str):
    """Add a bold label with normal value on the same line."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.name = 'Calibri'
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run_value = p.add_run(value)
    run_value.font.name = 'Calibri'
    run_value.font.size = Pt(11)
    run_value.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def markdown_to_docx(markdown_text: str, use_case_name: str = "Incident Response Plan") -> io.BytesIO:
    """Convert the AI-generated markdown response plan into a formatted DOCX."""
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    setup_styles(doc)
    
    lines = markdown_text.strip().split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Main title (## Incident Response Plan: ...)
        if stripped.startswith('## Incident Response Plan:') or stripped.startswith('## Incident Response Plan —'):
            title_text = stripped.lstrip('#').strip()
            p = doc.add_paragraph(title_text, style='IR Title')
            # Add a line separator after title
            p_sep = doc.add_paragraph()
            pPr = p_sep._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): '1B3A5C'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Section headers (### SECTION 1, ### SECTION 2, ### SECTION 3)
        if stripped.startswith('### '):
            heading_text = stripped.lstrip('#').strip()
            doc.add_heading(heading_text, level=1)
            i += 1
            continue
        
        # Sub-headers (#### or **Step N — ...)
        if stripped.startswith('#### '):
            heading_text = stripped.lstrip('#').strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue
        
        # Horizontal rules
        if stripped == '---' or stripped == '***':
            p_sep = doc.add_paragraph()
            pPr = p_sep._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '4',
                qn('w:space'): '1',
                qn('w:color'): 'CCCCCC'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Metadata fields (**Use Case:** ..., **Severity:** ...)
        metadata_match = re.match(r'^\*\*(.+?):\*\*\s*(.*)', stripped)
        if metadata_match and not stripped.startswith('**→') and not stripped.startswith('**Indicators'):
            label = metadata_match.group(1)
            value = metadata_match.group(2)
            add_metadata_field(doc, label, value)
            i += 1
            continue
        
        # Team headings (**→ EDR/CrowdStrike Team:** etc.)
        team_match = re.match(r'^\*\*→\s*(.+?):\*\*', stripped)
        if team_match:
            team_name = "→ " + team_match.group(1)
            p = doc.add_paragraph()
            p.style = doc.styles['Team Heading']
            run = p.add_run(team_name)
            i += 1
            continue
        
        # Bold sub-headers like **Indicators Pointing Toward TRUE POSITIVE:**
        bold_header_match = re.match(r'^\*\*(.+?)\*\*\s*$', stripped)
        if bold_header_match:
            doc.add_heading(bold_header_match.group(1), level=2)
            i += 1
            continue
        
        # Warning/log gap lines
        if stripped.startswith('⚠️'):
            p = doc.add_paragraph()
            p.style = doc.styles['Warning Text']
            # Parse bold parts within warning
            add_formatted_runs(p, stripped)
            i += 1
            continue
        
        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_runs(p, bullet_text)
            i += 1
            continue
        
        # Step headers (bold text like **Step 1 — ...**)
        step_match = re.match(r'^\*\*(Step \d+.+?)\*\*', stripped)
        if step_match:
            doc.add_heading(step_match.group(1), level=3)
            i += 1
            continue
        
        # Regular paragraph with possible inline formatting
        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def add_formatted_runs(paragraph, text: str):
    """Parse inline markdown formatting and add runs with appropriate styling."""
    # Pattern to match **bold**, *italic*, and `code` inline
    pattern = r'(\*\*(.+?)\*\*|`(.+?)`|\*(.+?)\*|([^*`]+))'
    
    for match in re.finditer(pattern, text):
        full = match.group(0)
        
        if full.startswith('**') and full.endswith('**'):
            # Bold text
            content = match.group(2)
            run = paragraph.add_run(content)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        elif full.startswith('`') and full.endswith('`'):
            # Inline code
            content = match.group(3)
            run = paragraph.add_run(content)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif full.startswith('*') and full.endswith('*') and not full.startswith('**'):
            # Italic text
            content = match.group(4)
            run = paragraph.add_run(content)
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        else:
            # Regular text
            content = match.group(5) if match.group(5) else full
            if content:
                run = paragraph.add_run(content)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
